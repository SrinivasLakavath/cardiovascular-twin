import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def markdown_to_docx(md_filepath, docx_filepath):
    print(f"Converting {md_filepath} to {docx_filepath}...")
    
    with open(md_filepath, 'r', encoding='utf-8') as f:
        md_text = f.read()
        
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = md_text.split('\n')
    
    in_code_block = False
    
    for line in lines:
        stripped_line = line.strip()
        
        # Handle Code Blocks (Mermaid / Scripts)
        if stripped_line.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                p = doc.add_paragraph()
                p.style = doc.styles['No Spacing']
                p.paragraph_format.left_indent = Pt(20)
                run = p.add_run("[CODE BLOCK / MERMAID DIAGRAM OMITTED FOR WORD - SEE MARKDOWN]")
                run.font.color.rgb = RGBColor(128, 128, 128)
                run.italic = True
            continue
            
        if in_code_block:
            continue
            
        if not stripped_line:
            continue

        # Headings
        if stripped_line.startswith('# '):
            p = doc.add_heading(stripped_line[2:], level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif stripped_line.startswith('## '):
            doc.add_heading(stripped_line[3:], level=2)
        elif stripped_line.startswith('### '):
            doc.add_heading(stripped_line[4:], level=3)
        elif stripped_line.startswith('#### '):
            doc.add_heading(stripped_line[5:], level=4)
            
        # Lists
        elif stripped_line.startswith('* ') or stripped_line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            # Basic bold parsing for lists
            text = stripped_line[2:]
            _add_formatted_runs(p, text)
            
        elif re.match(r'^\d+\.\s', stripped_line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s', '', stripped_line)
            _add_formatted_runs(p, text)
            
        # Horizontal Rule
        elif stripped_line.startswith('---'):
            doc.add_paragraph("_" * 50)
            
        # Standard Paragraphs
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, stripped_line)

    doc.save(docx_filepath)
    print(f"Successfully saved {docx_filepath}")

def _add_formatted_runs(paragraph, text):
    """Deepest respect for basic markdown bold/italic inside lines"""
    # Simply strip out massive image tags for word
    text = re.sub(r'!\[.*?\]\(.*?\)', '[IMAGE]', text)
    # Strip basic math formatting $...$
    text = re.sub(r'\$(.*?)\$', r'\1', text)
    
    # Very basic bold splitting (not perfect AST but works for this direct conversion)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Check for italics
            subparts = re.split(r'(\*.*?\*)', part)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*'):
                    run = paragraph.add_run(subpart[1:-1])
                    run.italic = True
                else:
                    # Strip inline code ticks
                    clean_text = subpart.replace('`', '')
                    paragraph.add_run(clean_text)

if __name__ == "__main__":
    docs_dir = os.path.join(os.path.dirname(__file__))
    
    report_md = os.path.join(docs_dir, "Project_Report.md")
    report_docx = os.path.join(docs_dir, "Project_Report.docx")
    
    guide_md = os.path.join(docs_dir, "Comprehensive_Guide.md")
    guide_docx = os.path.join(docs_dir, "Comprehensive_Guide.docx")
    
    if os.path.exists(report_md):
        markdown_to_docx(report_md, report_docx)
    
    if os.path.exists(guide_md):
        markdown_to_docx(guide_md, guide_docx)
